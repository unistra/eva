from django.db.models import Q

from docx import Document
from docx.shared import Cm
from bs4 import BeautifulSoup as bs

from mecc.apps.rules.models import Rule, Paragraph
from mecc.apps.training.models import Training, SpecificParagraph, AdditionalParagraph

def docx_gen(data):
    """
    Microsoft Word docx document generation for the rules
    """

    # ########## Some useful functions
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


    def add_lines(nb_lines=1):
        while nb_lines > 0:
            doc.add_paragraph()
            nb_lines -= 1


    def clean_up(text):
        """
        Paragraphs are stored as HTML paragraphs (with tags like <p></p>)
        This function parses paragraphs and returns a list that contains
        style elements and the text to write
        Ex : ['p','some text to write', 'i','some text to write in italic']
        """
        text = text.replace('\r\n', '').replace('\t', '')
        soup = bs(text, 'lxml')

        paragraphs = soup.body
        to_write = []

        for paragraph in paragraphs.descendants:
            if paragraph.name:
                to_write.append(paragraph.name)
            else:
                to_write.append(paragraph)
        if "ul" in to_write:
            to_write.remove("ul")

        return to_write


    # ########## write_ functions (used to write things in the document)
    def write_doc_subtitle(year, institute):
        doc.add_paragraph(year, style='Subtitle')
        doc.add_paragraph(institute, style='Subtitle')
        add_lines(5)


    def write_degree_label(label):
        doc.add_paragraph(label.upper(), style=style_title1)
        add_lines()
    

    def write_rules_type_label(rules_type):
        if rules_type is "standard":
            doc.add_paragraph(
                "Règles standards applicables à TOUS les diplômes "+\
                element['degree_type'].short_label.upper(), 
                style=style_title2
            )
        elif rules_type is "ccct":
            doc.add_paragraph(
                "Règles standards applicables aux diplômes "+\
                element['degree_type'].short_label.upper()+\
                " en régime CC/CT", 
                style=style_title2
            )
        elif rules_type is "eci":
            doc.add_paragraph(
                "Règles standards applicables aux diplômes "+\
                element['degree_type'].short_label.upper()+\
                " en régime ECI ",
                style=style_title2
            )
        else:
            pass

        add_lines()


    def write_rules_paragraphs(paragraphs):
        char_styles = []
        for paragraph in paragraphs:
            if isinstance(paragraph, Paragraph):
                to_write = clean_up(paragraph.text_standard)
            elif isinstance(paragraph, SpecificParagraph):
                to_write = clean_up(paragraph.text_specific_paragraph)
            else:
                to_write = clean_up(paragraph.text_additional_paragraph)
            for element in to_write:
                if element in ['i', 'em', 'strong', 'u', 'li']:
                    char_styles.append(element)
                elif element is not 'p':
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element)
                    for style in char_styles:
                        if style == 'i' or style == 'em':
                            run.italic = True
                        if style == 'strong':
                            run.bold = True
                        if style == 'u':
                            run.underline = True
                        if style == 'li':
                            paragraph.style = style_list_bullet
                    char_styles.clear()
            add_lines()


    def write_degree_rules(rules):
        
        for rule in rules:
            doc.add_paragraph(rule.label, style=style_title3)
            rule_paragraphs = Paragraph.objects.filter(rule=rule).order_by('display_order')
            write_rules_paragraphs(rule_paragraphs)


    def write_training_header(training, reference):
        doc.add_paragraph(training.label, style=style_title4)
        training_infos = doc.add_paragraph()
        training_infos.paragraph_format.tab_stops.add_tab_stop(Cm(6))
        training_infos.add_run(
            "%s - %s\t\t" % (training.get_MECC_type_display(), training.get_session_type_display())).bold = True
        if reference == "with_si":
            training_infos.add_run("Référence APOGEE : ").bold = True
            training_infos.add_run(training.ref_si_scol)
        if reference == "with_rof":
            training_infos.add_run("Référence ROF : ").bold = True
            training_infos.add_run(training.ref_cpa_rof)
        training_infos.style = style_title5

        in_charge = doc.add_paragraph()
        in_charge.add_run("Responsable(s) : ").bold = True
        in_charge.add_run(', '.join(respform for respform in training.get_respform_names))


    def write_training_specifics_additionals(rules, specifics, additionals):
        for rule in rules:
            doc.add_paragraph(rule.label, style=style_title3)
            if specifics.filter(rule_gen_id=rule.id):
                specifics_title = doc.add_paragraph()
                run = specifics_title.add_run("Dérogations :")
                run.bold = run.italic = True
                write_rules_paragraphs(specifics.filter(rule_gen_id=rule.id))
            if additionals.filter(rule_gen_id=rule.id):
                additionals_title = doc.add_paragraph()
                run = additionals_title.add_run("Alinéa additionnel :")
                run.bold = run.italic = True
                write_rules_paragraphs(additionals.filter(rule_gen_id=rule.id))


    # ########## Treatment depends on the document model
    model = data.pop(0)
    reference = data.pop(0)

    # ########## The generated document is based on a template
    # ########## The template contains all the styles we will use
    doc = Document('mecc/static/docx/eva-template.docx')

    # ########## Template clean up : we only want the title section
    doc_paragraphs = doc.paragraphs[2:]
    for paragraph in doc_paragraphs:
        delete_paragraph(paragraph)
    add_lines(2)

    # ########## Subtitle section (university year and cmp label)
    year = data.pop(0)
    code_year = year.code_year
    label_year = year.label_year
    institute = data.pop(0)
    write_doc_subtitle(label_year, institute)

    # ########## Initialization of the titles styles
    styles = doc.styles
    for style in styles:
        if 'Heading 1' in style.name:
            style_title1 = style
        if 'Heading 2' in style.name:
            style_title2 = style
        if 'Heading 3' in style.name:
            style_title3 = style
        if 'Heading 4' in style.name:
            style_title4 = style
        if 'Heading 5' in style.name:
            style_title5 = style
        if 'List Bullet' in style.name:
            style_list_bullet = style

    if model is 'a':

        for element in data:

            degree_rules = Rule.objects.filter(
                code_year=code_year,
                degree_type=element['degree_type']
            ).order_by('display_order')
            standard_rules = degree_rules.filter(is_ccct=1, is_eci=1)
            ccct_rules = degree_rules.filter(is_ccct=1, is_eci=0)
            eci_rules = degree_rules.filter(is_ccct=0, is_eci=1)

            write_degree_label(element['degree_type'].short_label)
            write_rules_type_label("standard")

            write_degree_rules(standard_rules)

            for mecc_type in element['mecc_types']:
                if mecc_type is 'C':
                    write_rules_type_label("ccct")
                    write_degree_rules(ccct_rules)
                elif mecc_type is 'E':
                    write_rules_type_label("eci")
                    write_degree_rules(eci_rules)

            degree_specifics = SpecificParagraph.objects.filter(
                training__in=element['trainings']
            )
            degree_additionals = AdditionalParagraph.objects.filter(
                training__in=element['trainings']
            )
            if degree_specifics or degree_additionals:
                doc.add_paragraph(
                    "Dérogations et alinéas additionnels",
                    style=style_title2
                )
                add_lines()

                for training in element['trainings']:
                    training_specifics = degree_specifics.filter(
                        training=training
                    )
                    training_additionals = degree_additionals.filter(
                        training=training
                    )
                    rules_with_specific_or_additional = degree_rules.filter(
                        Q(id__in=[specific.rule_gen_id for specific in training_specifics]) \
                        | \
                        Q(id__in=[additional.rule_gen_id for additional in training_additionals])
                    ).order_by('display_order')
                    write_training_header(training, reference)
                    write_training_specifics_additionals(
                        rules_with_specific_or_additional,
                        training_specifics,
                        training_additionals
                    )
            doc.add_page_break()
    else:
        
        for element in data:
            rules = Rule.objects.filter(
                code_year=code_year,
                degree_type=element['degree_type']
            ).order_by('display_order')
            
            write_degree_label(element['degree_type'].short_label)

            for training in element['trainings']:
                write_training_header(training, reference)
                doc.add_paragraph(
                    "Règles applicables à la formation",
                    style=style_title2
                )
                add_lines()
                if training.MECC_type == 'C':
                    training_rules = rules.exclude(is_eci=0)
                if training.MECC_type == 'E':
                    training_rules = rules.exclude(is_ccct=0)
                training_specifics = SpecificParagraph.objects.filter(training=training)
                training_additionals = AdditionalParagraph.objects.filter(training=training)
                for rule in training_rules:
                    doc.add_paragraph(rule.label, style=style_title3)
                    paragraphs = Paragraph.objects.filter(rule=rule).order_by('display_order')
                    to_write = []
                    for paragraph in paragraphs:
                        if paragraph.id in [specific.paragraph_gen_id for specific in training_specifics]:
                            to_write.append(training_specifics.get(
                                paragraph_gen_id=paragraph.id
                            )) 
                        else:
                            to_write.append(paragraph)
                    write_rules_paragraphs(to_write)
                    if rule.id in [additional.rule_gen_id for additional in training_additionals]:
                        to_write = training_additionals.filter(
                            rule_gen_id=rule.id
                        )
                        write_rules_paragraphs(to_write)
            doc.add_page_break()

    return doc
