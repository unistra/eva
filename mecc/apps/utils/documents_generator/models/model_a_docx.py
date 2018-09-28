from django.db.models import Q
from django.http import HttpResponse

from docx import Document
from docx.shared import Cm
from bs4 import BeautifulSoup as bs

from mecc.apps.institute.models import Institute
from mecc.apps.rules.models import Rule, Paragraph
from mecc.apps.training.models import Training, SpecificParagraph, AdditionalParagraph
from mecc.apps.utils.queries import currentyear
from mecc.apps.years.models import UniversityYear

from ..document import Document as Doc

class ModelADocx(Doc):
    def __init__(self, reference, trainings):

        # ########## Treatment depends on the document model
        self.model = 'a'
        self.reference = reference
        self.trainings = trainings.order_by(
            'degree_type',
            'MECC_type',
            'label'
        ).select_related('degree_type')
        self.code_year = currentyear().code_year

        # ########## HTTP response init
        doc_name = "eva_rules_%s" % self.code_year
        self.response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.response['Content-Disposition'] = 'attachment; filename="%s".docx' % doc_name

        # ########## The generated document is based on a template
        # ########## The template contains all the styles we will use
        self.doc = Document('mecc/static/docx/eva-template.docx')



        # ########## Subtitle section (university year and cmp label)
        self.year = UniversityYear.objects.get(
            code_year=currentyear().code_year
        ).label_year
        self.institute = Institute.objects.get(
            code=trainings.first().supply_cmp
        ).label

        # ########## Initialization of the titles styles
        styles = self.doc.styles
        for style in styles:
            if 'Heading 1' in style.name:
                self.style_title1 = style
            if 'Heading 2' in style.name:
                self.style_title2 = style
            if 'Heading 3' in style.name:
                self.style_title3 = style
            if 'Heading 4' in style.name:
                self.style_title4 = style
            if 'Heading 5' in style.name:
                self.style_title5 = style
            if 'List Bullet' in style.name:
                self.style_list_bullet = style

    def build_doc(self):

        # ########## Template clean up : we only want the title section
        doc_paragraphs = self.doc.paragraphs[2:]
        for paragraph in doc_paragraphs:
            self.delete_paragraph(paragraph)
        self.add_lines(2)

        self.write_doc_subtitle()

        degree_types = list(set([training.degree_type for training in self.trainings]))

        for degree_type in degree_types:
            degree_rules = Rule.objects.filter(
                code_year=self.code_year,
                degree_type=degree_type
            ).order_by('display_order')
            standard_rules = degree_rules.filter(is_ccct=1, is_eci=1)

            self.write_degree_label(degree_type.short_label)
            self.write_rules_type_label(
                "standard",
                degree_type.short_label
            )
            self.write_degree_rules(standard_rules)

            mecc_types = [training.MECC_type for training in self.trainings.filter(
                degree_type=degree_type
            )]
            if 'C' in mecc_types:
                ccct_rules = degree_rules.filter(is_ccct=1, is_eci=0)
                self.write_rules_type_label(
                    "ccct",
                    degree_type.short_label
                )
                self.write_degree_rules(ccct_rules)
            if 'E' in mecc_types:
                eci_rules = degree_rules.filter(is_ccct=0, is_eci=1)
                self.write_rules_type_label(
                    "eci",
                    degree_type.short_label
                )
                self.write_degree_rules(eci_rules)

            degree_specifics = SpecificParagraph.objects.filter(
                training__in=self.trainings.filter(degree_type=degree_type)
            )
            degree_additionals = AdditionalParagraph.objects.filter(
                training__in=self.trainings.filter(degree_type=degree_type)
            )
            if degree_specifics or degree_additionals:
                self.doc.add_paragraph(
                    "Dérogations et alinéas additionnels",
                    style=self.style_title2
                )
                self.add_lines()

                for training in self.trainings.filter(degree_type=degree_type):
                    training_specifics = degree_specifics.filter(
                        training=training
                    )
                    training_additionals = degree_additionals.filter(
                        training=training
                    )
                    rules_with_specific_or_additional = degree_rules.filter(
                        Q(id__in=[specific.rule_gen_id for specific in training_specifics]) \
                        | \
                        Q(id__in=[additional.rule_gen_id for additional in training_additionals])
                    ).order_by('display_order')
                    self.write_training_header(training)
                    self.write_training_specifics_additionals(
                        rules_with_specific_or_additional,
                        training_specifics,
                        training_additionals
                    )
            self.doc.add_page_break()

        self.doc.save(self.response)

        return self.response

    # ########## Some useful functions
    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def add_lines(self, nb_lines=1):
        while nb_lines > 0:
            self.doc.add_paragraph()
            nb_lines -= 1

    @staticmethod
    def clean_up(text):
        """
        Paragraphs are stored as HTML paragraphs (with tags like <p></p>)
        This function parses paragraphs and returns a list that contains
        style elements and the text to write
        Ex : ['p','some text to write', 'i','some text to write in italic']
        """
        text = text.replace('\r\n', '').replace('\t', '')
        soup = bs(text, 'lxml')

        paragraphs = soup.body
        to_write = []

        for paragraph in paragraphs.descendants:
            if paragraph.name:
                to_write.append(paragraph.name)
            else:
                to_write.append(paragraph)
        if "ul" in to_write:
            to_write.remove("ul")

        return to_write

    # ########## write_ functions (used to write things in the document)
    def write_doc_subtitle(self):
        self.doc.add_paragraph(self.year, style='Subtitle')
        self.doc.add_paragraph(self.institute, style='Subtitle')
        self.add_lines(5)

    def write_degree_label(self, label):
        self.doc.add_paragraph(label.upper(), style=self.style_title1)
        self.add_lines()


    def write_rules_type_label(self, rules_type, degree_type):
        if rules_type == "standard":
            self.doc.add_paragraph(
                "Règles standards applicables à TOUS les diplômes "+\
                degree_type.upper(),
                style=self.style_title2
            )
        if rules_type == "ccct":
            self.doc.add_paragraph(
                "Règles standards applicables aux diplômes "+\
                degree_type.upper()+\
                " en régime CC/CT",
                style=self.style_title2
            )
        if rules_type == "eci":
            self.doc.add_paragraph(
                "Règles standards applicables aux diplômes "+\
                degree_type.upper()+\
                " en régime ECI ",
                style=self.style_title2
            )

        self.add_lines()

    def write_degree_rules(self, rules):

        for rule in rules:
            self.doc.add_paragraph(rule.label, style=self.style_title3)
            rule_paragraphs = Paragraph.objects.filter(rule=rule).order_by('display_order')
            self.write_rules_paragraphs(rule_paragraphs)

    def write_rules_paragraphs(self, paragraphs):
        char_styles = []
        for paragraph in paragraphs:
            if isinstance(paragraph, Paragraph):
                to_write = self.clean_up(paragraph.text_standard)
            elif isinstance(paragraph, SpecificParagraph):
                to_write = self.clean_up(paragraph.text_specific_paragraph)
            else:
                to_write = self.clean_up(paragraph.text_additional_paragraph)
            for element in to_write:
                if element in ['i', 'em', 'strong', 'u', 'li']:
                    char_styles.append(element)
                elif element != 'p':
                    paragraph = self.doc.add_paragraph()
                    run = paragraph.add_run(element)
                    for style in char_styles:
                        if style == 'i' or style == 'em':
                            run.italic = True
                        if style == 'strong':
                            run.bold = True
                        if style == 'u':
                            run.underline = True
                        if style == 'li':
                            paragraph.style = self.style_list_bullet
                    char_styles.clear()
            self.add_lines()

    def write_training_header(self, training):
        self.doc.add_paragraph(training.label, style=self.style_title4)
        training_infos = self.doc.add_paragraph()
        training_infos.paragraph_format.tab_stops.add_tab_stop(Cm(6))
        training_infos.add_run(
            "%s - %s\t\t" % (training.get_MECC_type_display(), training.get_session_type_display())).bold = True
        if self.reference == "with_si":
            training_infos.add_run("Référence APOGEE : ").bold = True
            training_infos.add_run(training.ref_si_scol)
        if self.reference == "with_rof":
            training_infos.add_run("Référence ROF : ").bold = True
            training_infos.add_run(training.ref_cpa_rof)
        training_infos.style = self.style_title5

        in_charge = self.doc.add_paragraph()
        in_charge.add_run("Responsable(s) : ").bold = True
        in_charge.add_run(', '.join(respform for respform in training.get_respform_names))

    def write_training_specifics_additionals(self, rules, specifics, additionals):
        for rule in rules:
            self.doc.add_paragraph(rule.label, style=self.style_title3)
            if specifics.filter(rule_gen_id=rule.id):
                specifics_title = self.doc.add_paragraph()
                run = specifics_title.add_run("Dérogations :")
                run.bold = run.italic = True
                self.write_rules_paragraphs(specifics.filter(rule_gen_id=rule.id))
            if additionals.filter(rule_gen_id=rule.id):
                additionals_title = self.doc.add_paragraph()
                run = additionals_title.add_run("Alinéa additionnel :")
                run.bold = run.italic = True
                self.write_rules_paragraphs(additionals.filter(rule_gen_id=rule.id))
